"""
Document Processor Module

Handles loading and processing of different document types with smart chunking.
Supports PDF, DOCX, and plain text files.
"""
import os
import re
import magic
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass

import spacy
from spacy.language import Language
from spacy.lang.en import English
from spacy.lang.fr import French
from spacy.lang.de import German
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    PythonCodeTextSplitter
)
from typing import List, Dict, Any, Optional, Tuple, Generator, Union
from dataclasses import dataclass, field
import re
import logging

logger = logging.getLogger(__name__)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

@dataclass
class DocumentChunk:
    """Represents a chunk of text with metadata."""
    text: str
    metadata: Dict[str, Any]

class DocumentProcessor:
    """Processes different document types and splits them into semantic chunks."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize the document processor with multi-language support and semantic chunking.
        
        Args:
            chunk_size: Maximum size of each text chunk (in characters)
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Load language models with medium-sized models for better accuracy
        self.nlp_models = {}
        try:
            self.nlp_models['en'] = spacy.load('en_core_web_md')
            # Use the default English model for backward compatibility
            self.nlp = self.nlp_models.get('en', spacy.blank('en'))
            logger.info("Loaded English language model (medium)")
        except OSError:
            logger.warning("English medium model not found, downloading...")
            os.system('python -m spacy download en_core_web_md')
            self.nlp_models['en'] = spacy.load('en_core_web_md')
            
        try:
            self.nlp_models['fr'] = spacy.load('fr_core_news_md')
            logger.info("Loaded French language model (medium)")
        except OSError:
            logger.warning("French medium model not found, downloading...")
            os.system('python -m spacy download fr_core_news_md')
            self.nlp_models['fr'] = spacy.load('fr_core_news_md')
            
        try:
            self.nlp_models['de'] = spacy.load('de_core_news_md')
            logger.info("Loaded German language model (medium)")
        except OSError as e:
            logger.warning("German medium model not found, downloading...")
            os.system('python -m spacy download de_core_news_md')
            self.nlp_models['de'] = spacy.load('de_core_news_md')
        
        # Set default language
        self.default_lang = 'en'
        
        # Configure text splitters
        self.base_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n\n", "\n\n", "\n", ". ", "? ", "! ", " ", ""]
        )
        
        # Semantic chunking configuration
        self.semantic_separators = [
            (re.compile(r'\n#{1,3}\s+'), 'heading'),  # Markdown headings
            (re.compile(r'\n\s*[A-Z][^\n]{20,}[\.!?]\s+'), 'sentence'),  # Sentence endings
            (re.compile(r'\n\s*\*\s+'), 'list_item'),  # List items
            (re.compile(r'\n\s*\d+\.\s+'), 'numbered_item'),  # Numbered lists
        ]
        
        # Semantic chunk size thresholds (in characters)
        self.min_chunk_size = chunk_size // 4
        self.max_chunk_size = chunk_size * 2
        
    def process_file(self, filepath: str) -> List[DocumentChunk]:
        """Process a file and return its chunks with metadata.
        
        Args:
            filepath: Path to the file to process
            
        Returns:
            List of DocumentChunk objects
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
            
        try:
            # Get file stats for metadata
            file_stats = os.stat(filepath)
            file_size = file_stats.st_size / (1024 * 1024)  # Size in MB
            
            # Determine file type
            file_type = self._detect_file_type(filepath)
            
            # Extract text based on file type
            if file_type == 'pdf':
                text = self._extract_pdf_text(filepath)
            elif file_type == 'docx':
                text = self._extract_docx_text(filepath)
            elif file_type == 'text':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not text.strip():
                logger.warning(f"No content extracted from {filepath}")
                return []
                
            # Detect language
            detected_lang = self._detect_language(text)
            self.nlp = self.nlp_models.get(detected_lang, self.nlp)
            
            # Add file metadata
            metadata = {
                'source': os.path.basename(filepath),
                'full_path': filepath,
                'file_type': file_type,
                'file_size_mb': round(file_size, 2),
                'last_modified': file_stats.st_mtime,
                'created': file_stats.st_ctime,
                'language': detected_lang
            }
            
            # Clean and chunk the text with metadata
            return self._chunk_text(text, metadata)
            
        except Exception as e:
            logger.error(f"Error processing file {filepath}: {str(e)}", exc_info=True)
            raise
    
    def _extract_docx_text(self, filepath: str) -> str:
        """Extract text from a DOCX file with formatting.
        
        Args:
            filepath: Path to the DOCX file
            
        Returns:
            Extracted text as a string with preserved structure
        """
        try:
            doc = DocxDocument(filepath)
            text_parts = []
            
            # Extract paragraphs with their styles
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                logger.warning(f"No extractable text found in DOCX: {filepath}")
                return ""
                
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {filepath}: {str(e)}", exc_info=True)
            raise

    def _detect_file_type(self, filepath: str) -> str:
        """Detect the type of a file with improved MIME type checking.
        
        Args:
            filepath: Path to the file
            
        Returns:
            File type as a string (pdf, docx, text, etc.)
        """
        try:
            # First check by extension for known types
            file_ext = os.path.splitext(filepath)[1].lower()
            
            if file_ext in ['.pdf']:
                return 'pdf'
            elif file_ext in ['.docx', '.doc']:
                return 'docx'
            elif file_ext in ['.txt', '.md', '.markdown']:
                return 'text'
                
            # Fall back to MIME type detection
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(filepath).lower()
            
            if 'pdf' in mime_type:
                return 'pdf'
            elif 'wordprocessingml' in mime_type or 'officedocument.word' in mime_type:
                return 'docx'
            elif 'text' in mime_type or filepath.lower().endswith(('.txt', '.md', '.markdown')):
                return 'text'
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
                
        except Exception as e:
            logger.error(f"Error detecting file type for {filepath}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, filepath: str) -> str:
        """Extract text from a PDF file with improved formatting.
        
        Args:
            filepath: Path to the PDF file
            
        Returns:
            Extracted text as a string with preserved structure
        """
        try:
            reader = PdfReader(filepath)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    # Extract text with layout information
                    page_text = page.extract_text()
                    
                    # Clean up the text
                    if page_text:
                        # Remove excessive whitespace but preserve paragraph breaks
                        page_text = re.sub(r'\s+', ' ', page_text).strip()
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                        
                except Exception as page_error:
                    logger.warning(f"Error processing page {page_num} of {filepath}: {str(page_error)}")
                    continue
            
            if not text_parts:
                logger.warning(f"No extractable text found in PDF: {filepath}")
                return ""
                
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {filepath}: {str(e)}", exc_info=True)
            raise
    
    def _chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into semantic chunks with enhanced metadata.
        
        This method uses a multi-level chunking strategy:
        1. Split by major semantic boundaries (headings, sections)
        2. Further split by paragraphs and sentences
        3. Ensure chunks respect semantic units and size constraints
        
        Args:
            text: Text to chunk
            metadata: Metadata dictionary for the source file
            
        Returns:
            List of DocumentChunk objects with semantic metadata
        """
        text = text.strip()
        if not text:
            return []
            
        # Enhanced metadata
        chunk_metadata = metadata.copy()
        chunks = []
        
        # First pass: Split by major semantic boundaries
        segments = self._split_by_semantic_boundaries(text)
        
        # Second pass: Process each segment
        for segment, segment_type in segments:
            if not segment.strip():
                continue
                
            # Update metadata with segment type
            segment_metadata = chunk_metadata.copy()
            segment_metadata['chunk_type'] = segment_type
            
            # If segment is small, add as is
            if len(segment) <= self.chunk_size + self.chunk_overlap:
                chunks.append(DocumentChunk(
                    text=segment,
                    metadata=segment_metadata
                ))
                continue
                
            # Otherwise, split further by paragraphs and sentences
            sub_chunks = self._split_by_paragraphs_and_sentences(segment, segment_metadata)
            chunks.extend(sub_chunks)
        
        # Final pass: Merge small chunks where possible
        chunks = self._merge_small_chunks(chunks)
        
        return chunks
        
    def _split_by_semantic_boundaries(self, text: str) -> List[tuple]:
        """Split text by semantic boundaries like headings, sections, etc."""
        segments = [(text, 'document')]  # Default to 'document' type
        
        # Try each semantic separator in order of priority
        for pattern, seg_type in self.semantic_separators:
            new_segments = []
            
            for segment, current_type in segments:
                # Skip if segment is already small enough
                if len(segment) <= self.chunk_size:
                    new_segments.append((segment, current_type))
                    continue
                    
                # Split by the current pattern
                parts = pattern.split(segment)
                if len(parts) <= 1:  # No splits occurred
                    new_segments.append((segment, current_type))
                    continue
                    
                # Rebuild segments with proper types
                for i, part in enumerate(parts):
                    if not part.strip():
                        continue
                        
                    # Every odd part is a separator (captured in split)
                    if i % 2 == 1:
                        new_segments.append((part, seg_type))
                    else:
                        new_segments.append((part, current_type))
            
            segments = new_segments
            
        return segments
        
    def _split_by_paragraphs_and_sentences(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into chunks respecting paragraphs and sentences."""
        chunks = []
        current_chunk = []
        current_length = 0
        
        # Split into paragraphs first
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        for para in paragraphs:
            # Process each paragraph with spaCy for sentence segmentation
            doc = self.nlp(para)
            sentences = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
            
            for sent in sentences:
                sent_length = len(sent)
                
                # If adding this sentence would exceed max chunk size, finalize current chunk
                if current_chunk and (current_length + sent_length > self.max_chunk_size):
                    chunk_text = ' '.join(current_chunk)
                    chunks.append(DocumentChunk(
                        text=chunk_text,
                        metadata=metadata.copy()
                    ))
                    current_chunk = []
                    current_length = 0
                
                current_chunk.append(sent)
                current_length += sent_length
        
        # Add the last chunk if there's any remaining text
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append(DocumentChunk(
                text=chunk_text,
                metadata=metadata.copy()
            ))
        
        return chunks
        
    def _merge_small_chunks(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Merge small chunks to avoid too many tiny chunks."""
        if not chunks:
            return []
            
        merged_chunks = []
        current_chunk = chunks[0].text
        current_metadata = chunks[0].metadata.copy()
        current_length = len(current_chunk)
        
        for chunk in chunks[1:]:
            chunk_text = chunk.text
            chunk_length = len(chunk_text)
            
            # If merging would keep us under chunk size, merge
            if current_length + chunk_length <= self.chunk_size + self.chunk_overlap:
                current_chunk += ' ' + chunk_text
                current_length += chunk_length + 1  # +1 for space
            else:
                # Otherwise, finalize current chunk and start a new one
                merged_chunks.append(DocumentChunk(
                    text=current_chunk,
                    metadata=current_metadata
                ))
                current_chunk = chunk_text
                current_metadata = chunk.metadata.copy()
                current_length = chunk_length
        
        # Add the last chunk
        if current_chunk:
            merged_chunks.append(DocumentChunk(
                text=current_chunk,
                metadata=current_metadata
            ))
            
        return merged_chunks
